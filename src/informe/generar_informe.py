"""
Entregable 1 — Informe Técnico (PDF), estructurado bajo las 4 fases.

Genera informe/Informe_Tecnico.docx con python-docx (con las figuras del EDA y
del modelo) y luego se convierte a PDF. Las cifras del informe se leen de los
JSON que produce el pipeline, así que si se vuelve a correr el ETL o el modelo,
el informe se regenera con los mismos números y no hay que retocarlo a mano.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RUTA_PROYECTO = Path(__file__).resolve().parents[2]
RUTA_REPORTES = RUTA_PROYECTO / "reports"
RUTA_FIGURAS = RUTA_REPORTES / "figures"
RUTA_PROCESSED = RUTA_PROYECTO / "data" / "processed"
RUTA_SALIDA = Path(
    os.environ.get("INFORME_SALIDA", RUTA_PROYECTO / "informe" / "Informe_Tecnico.docx")
)

VERDE = RGBColor(0x2E, 0x7D, 0x32)
GRIS = RGBColor(0x42, 0x42, 0x42)

INTEGRANTES = [
    "Chile Silva Carlos Junior",
    "Figueroa Arteaga Luis Emilio",
    "Menoscal Santana Bryan Steven",
]
DOCENTE = "Ing. Carlos Manosalvas G."

with open(RUTA_REPORTES / "eda_resultados.json", encoding="utf-8") as f:
    EDA = json.load(f)
with open(RUTA_REPORTES / "modelo_resultados.json", encoding="utf-8") as f:
    MODELO = json.load(f)
with open(RUTA_PROCESSED / "reporte_calidad.json", encoding="utf-8") as f:
    CALIDAD = json.load(f)

RECOMENDACIONES = pd.read_csv(RUTA_REPORTES / "recomendaciones_compra.csv")
FACT_REND = pd.read_csv(RUTA_PROCESSED / "fact_rendimiento.csv")
DIM_GEO = pd.read_csv(RUTA_PROCESSED / "dim_geografia.csv")
FACT_COMPRAS = pd.read_csv(RUTA_PROCESSED / "fact_compras.csv")

N_FILAS = len(FACT_REND)
N_PAISES = len(DIM_GEO)
N_COMPRAS = len(FACT_COMPRAS)
ANIO_MIN = int(FACT_REND["id_tiempo"].min())
ANIO_MAX = int(FACT_REND["id_tiempo"].max())

CONTEO_DECISIONES = RECOMENDACIONES["decision"].value_counts()
CULTIVOS_AUMENTAR = ", ".join(
    RECOMENDACIONES.loc[RECOMENDACIONES["decision"] == "AUMENTAR compra", "cultivo"].str.lower()
)

doc = Document()

estilo_normal = doc.styles["Normal"]
estilo_normal.font.name = "Calibri"
estilo_normal.font.size = Pt(11)


def titulo(texto: str, nivel: int = 1) -> None:
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = VERDE if nivel <= 2 else GRIS


def parrafo(texto: str, negrita: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def vinetas(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def figura(nombre_archivo: str, leyenda: str, ancho: float = 6.3) -> None:
    ruta = RUTA_FIGURAS / nombre_archivo
    doc.add_picture(str(ruta), width=Inches(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(leyenda)
    run.italic = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def tabla(encabezados: list[str], filas: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 6"
    for i, enc in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = enc
        for p in celda.paragraphs:
            for r in p.runs:
                r.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = str(valor)
    doc.add_paragraph()


# ============================ PORTADA =======================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\nUNIVERSIDAD LAICA ELOY ALFARO DE MANABÍ\n")
r.bold = True
r.font.size = Pt(16)
r = p.add_run("\nProyecto Integrador — Analítica de Negocios\nNivel 7A\n\n")
r.font.size = Pt(13)
r = p.add_run(
    "\nSistema de predicción del rendimiento\nde cosechas para decisiones de compra\n"
)
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = VERDE
r = p.add_run("\nAgroComercial del Litoral S.A. (caso simulado)\n")
r.italic = True
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\nDocente:\n")
r.bold = True
r.font.size = Pt(12)
r = p.add_run(f"{DOCENTE}\n\n")
r.font.size = Pt(12)
r = p.add_run("Integrantes:\n")
r.bold = True
r.font.size = Pt(12)
for nombre in INTEGRANTES:
    r = p.add_run(f"{nombre}\n")
    r.font.size = Pt(12)
r = p.add_run("\n\nManta — Manabí, Ecuador\n")
r.font.size = Pt(11)
r.font.color.rgb = GRIS
doc.add_page_break()

# ======================= RESUMEN EJECUTIVO ==================================
titulo("Resumen Ejecutivo", 1)
kpis = EDA["kpis"]
rf = MODELO["metricas_por_modelo"]["Random Forest"]
parrafo(
    "AgroComercial del Litoral S.A. es una comercializadora que compra cosechas a productores "
    "ecuatorianos y las revende a industrias, supermercados y compradores de exportación. Su "
    "ganancia depende de acertar cuánto comprar cada temporada. Hoy esa decisión se toma más por "
    "intuición que por datos, y eso deja a la empresa expuesta a dos errores caros: comprar de más "
    "y perder producto por deterioro, o comprar de menos y quedarse sin qué vender cuando hay "
    "demanda. Para atacar ese problema desarrollamos una solución analítica que va desde la "
    "preparación de los datos hasta un modelo predictivo funcional y un dashboard de apoyo a la "
    "decisión, y que estima el rendimiento esperado de los cultivos (ton/ha) para traducirlo en una "
    "recomendación de compra por cultivo."
)
parrafo(
    f"En números, integramos tres fuentes: los datos oficiales de rendimiento y pesticidas de "
    f"FAOSTAT, el clima obtenido en vivo desde la API pública del Banco Mundial y un registro "
    f"interno de compras (anonimizado). El resultado es un modelo dimensional en estrella con "
    f"{N_FILAS:,} observaciones limpias de {N_PAISES} países, entre {ANIO_MIN} y {ANIO_MAX}. De los "
    f"tres algoritmos que probamos, el mejor fue Random Forest, con un R² de {rf['r2']:.3f} y un "
    f"error de {rf['rmse_ton_ha']} ton/ha sobre datos que no vio en el entrenamiento. Sobre esa "
    f"predicción, la regla de compra sugiere aumentar el volumen en "
    f"{CONTEO_DECISIONES.get('AUMENTAR compra', 0)} cultivos ({CULTIVOS_AUMENTAR}), mantenerlo en "
    f"{CONTEO_DECISIONES.get('MANTENER volumen', 0)} y reducirlo en "
    f"{CONTEO_DECISIONES.get('REDUCIR compra', 0)} para el ciclo "
    f"{int(RECOMENDACIONES['anio_objetivo'].iloc[0])}."
)
doc.add_page_break()

# ============================== FASE 1 ======================================
titulo("Fase 1. Definición del Negocio y Entendimiento del Problema", 1)

titulo("1.1 Contexto del Negocio", 2)
parrafo(
    "AgroComercial del Litoral S.A. es una empresa simulada que usamos como caso de estudio. "
    "Funciona como intermediaria en la cadena agrícola: le compra la cosecha al productor, la "
    "almacena y la revende a industrias procesadoras, supermercados y mercados de exportación. Su "
    "margen sale de comprar el volumen adecuado, a buen precio y en el momento oportuno. El problema "
    "es que la materia prima que maneja depende de algo que la empresa no controla. El rendimiento "
    "de un cultivo cambia con el clima y con el manejo (riego, pesticidas), y eso hace que la oferta "
    "disponible varíe bastante de una temporada a otra."
)

titulo("1.2 Planteamiento del Problema (Cultura Data-Driven)", 2)
parrafo(
    "Hoy las compras se deciden de forma reactiva. El jefe de compras se guía por su experiencia y "
    "por lo que pasó la temporada anterior, sin ninguna estimación de cuánto van a rendir los "
    "cultivos con el clima que se espera. Ese es el dolor del negocio: no poder anticipar el "
    "rendimiento por país o región y temporada. Y cuando no se puede anticipar, aparecen dos "
    "situaciones que cuestan dinero:"
)
vinetas(
    [
        "Sobrestock: comprar más de lo que el mercado alcanza a absorber. El producto se almacena de "
        "más, se deteriora y termina en pérdida. En el registro interno la merma promedio por orden "
        f"es de {kpis['KPI_4_merma_promedio_pct']['valor']} % (dato del escenario simulado).",
        "Quiebre de stock: quedarse corto justo cuando hay demanda, con la venta perdida y el riesgo "
        "de perder también al cliente.",
    ]
)
parrafo(
    "Pasar a una cultura data-driven significa, en la práctica, dejar de decidir por corazonada y "
    "empezar a decidir con evidencia: usar el histórico de rendimiento, clima y pesticidas para "
    "estimar el rendimiento esperado y planificar las compras a partir de esa estimación. Que el "
    "riesgo existe no es una opinión, se puede medir: la variabilidad interanual del rendimiento en "
    f"Ecuador tiene un coeficiente de variación de {kpis['KPI_2_variabilidad_rendimiento_pct']['valor']} %, "
    "calculado directamente sobre los datos."
)

titulo("1.3 Objetivos de Analítica (los 4 pilares)", 2)
parrafo(
    "Definimos el alcance separando las cuatro preguntas que la analítica puede responder. Cada "
    "pilar tiene un objetivo concreto dentro de este proyecto:"
)
tabla(
    ["Pilar", "Pregunta", "Qué hacemos en este proyecto"],
    [
        [
            "Descriptiva",
            "¿Qué ha pasado?",
            "Caracterizar el rendimiento histórico por cultivo, país y año, y ver qué cultivos y "
            "zonas rinden más y cómo se han movido el clima y los pesticidas.",
        ],
        [
            "Diagnóstica",
            "¿Por qué pasó?",
            "Estudiar la relación entre el rendimiento y sus factores (lluvia, temperatura, "
            "pesticidas), cuidando de no confundir correlación con causalidad.",
        ],
        [
            "Predictiva",
            "¿Qué pasará?",
            "Estimar el rendimiento (ton/ha) esperado para un cultivo según las condiciones de clima "
            "y manejo, con modelos de regresión.",
        ],
        [
            "Prescriptiva",
            "¿Qué conviene hacer?",
            "Recomendar cuánto comprar y qué cultivos priorizar según el rendimiento predicho, con "
            "una regla de decisión que reduzca el sobrestock y el quiebre.",
        ],
    ],
)

titulo("1.4 Fuentes de Datos y Ética Inicial", 2)
parrafo(
    "Trabajamos con tres orígenes de datos, uno de ellos una API pública, tal como pedía la "
    "consigna:"
)
tabla(
    ["Tipo", "Fuente", "Contenido"],
    [
        [
            "Dataset (CSV)",
            "FAOSTAT — descarga oficial (bulks-faostat.fao.org, dominios QCL y RP)",
            f"Rendimiento (kg/ha) y pesticidas (t) por país, cultivo y año: {N_FILAS:,} filas, "
            f"{N_PAISES} países, 10 cultivos, {ANIO_MIN}-{ANIO_MAX}. Reemplaza al dataset de Kaggle "
            "que se había propuesto en la Fase 1, que solo llegaba hasta 2013.",
        ],
        [
            "API pública",
            "Banco Mundial: Climate Change Knowledge Portal (cckpapi.worldbank.org) y World Bank "
            "Open Data (api.worldbank.org/v2)",
            "Clima oficial por país y año, extraído en vivo: temperatura media (°C) y precipitación "
            "anual (mm) de la serie CRU TS 4.08 hasta 2023, más algunos indicadores agrícolas de "
            "referencia para Ecuador y cuatro países vecinos.",
        ],
        [
            "Excel interno (simulado)",
            "Registro de compras de la empresa",
            f"{N_COMPRAS:,} órdenes históricas con volúmenes, precios de compra y de reventa, canal, "
            "región y merma.",
        ],
    ],
)
parrafo("Sobre gobernanza y ética (privacidad desde el diseño):", negrita=True)
vinetas(
    [
        "Los datos de rendimiento y clima vienen agregados por país, así que no contienen datos "
        "personales. En el Excel de compras los proveedores aparecen solo con un código anónimo "
        "(PROV-001, PROV-002, y así), sin nombres, RUC, teléfonos ni ubicaciones exactas.",
        "Cada fuente y cada transformación queda registrada en el pipeline y en un reporte de "
        "calidad que se genera automáticamente (lo detallamos en la Fase 2), de modo que el origen "
        "de un número siempre se puede rastrear.",
        "El modelo se pensó como apoyo a la decisión, no como reemplazo del jefe de compras. Por eso "
        "las predicciones siempre se comunican junto con su margen de error.",
        "Somos conscientes de que el histórico sobrerrepresenta a ciertos países y cultivos. Ese "
        "sesgo lo medimos con números en la Fase 4.2 en lugar de darlo por descontado.",
    ]
)
doc.add_page_break()

# ============================== FASE 2 ======================================
titulo("Fase 2. Ingeniería de Datos y Modelado Multidimensional", 1)

titulo("2.1 Modelado Multidimensional (Esquema Estrella)", 2)
parrafo(
    "Diseñamos un esquema en estrella con dos tablas de hechos y cuatro dimensiones que ambas "
    "comparten (dimensiones conformadas). Optamos por estrella y no por copo de nieve porque nos "
    "interesaba que las consultas en Power BI fueran simples y con pocos joins. Las dimensiones son "
    "chicas, así que normalizarlas más no ahorraría casi nada y sí complicaría el modelo."
)
parrafo("Tablas de hechos:", negrita=True)
vinetas(
    [
        f"fact_rendimiento ({N_FILAS:,} filas). El grano es país por cultivo por año. Métricas: "
        "rendimiento_ton_ha, precipitacion_mm, pesticidas_ton y temperatura_media_c.",
        f"fact_compras ({N_COMPRAS:,} filas). El grano es la orden de compra. Métricas: "
        "volumen_comprado_ton, precio_compra_usd_ton, precio_reventa_usd_ton, costo_total_usd, "
        "ingreso_reventa_usd, margen_bruto_usd y merma_pct.",
    ]
)
parrafo("Tablas de dimensiones:", negrita=True)
vinetas(
    [
        f"dim_tiempo ({ANIO_MAX - ANIO_MIN + 1} filas): id_tiempo, anio, decada, periodo.",
        f"dim_geografia ({N_PAISES} filas): id_geografia, pais, region_comercial (Ecuador, "
        "Latinoamérica de referencia o resto del mundo).",
        "dim_cultivo (10 filas): id_cultivo, cultivo (nombre FAO en inglés), cultivo_es y categoria "
        "(cereal, tubérculo, oleaginosa o musácea).",
        "dim_proveedor (40 filas): id_proveedor, codigo_proveedor (anónimo) y region_origen (Costa, "
        "Sierra o Amazonía).",
    ]
)

titulo("2.2 Pipeline de Integración (ETL)", 2)
duplicados = CALIDAD["pasos"][0].get("duplicados_eliminados", 0)
parrafo(
    "El pipeline está hecho en Python (pandas y requests) y se divide en cuatro scripts que se "
    "pueden volver a correr sin problema. extraer_datos_actualizados.py une el bulk oficial de "
    "FAOSTAT con el clima de la API del Banco Mundial (guardamos una copia en caché local y "
    "reintentamos si la API falla por un momento). extract_api.py trae indicadores de referencia "
    "del Banco Mundial. generar_compras_internas.py arma la fuente interna simulada, ya sin datos "
    "personales desde el origen. Y transform_load.py se encarga de la limpieza, el tipado, la "
    "validación y la carga del esquema estrella. Cada corrida deja un reporte de calidad en JSON con "
    "lo que se hizo y su efecto."
)
parrafo("Calidad de datos (valores de la última corrida):", negrita=True)
vinetas(
    [
        "Integración por clave país-año. Homologamos los países a códigos ISO3 (y dejamos fuera a "
        "propósito los agregados regionales de FAOSTAT como 'Americas' o 'World'), y solo nos "
        "quedamos con las combinaciones país-cultivo-año que tienen las cuatro métricas completas.",
        f"Duplicados exactos eliminados: {duplicados:,} filas en la corrida final.",
        "Tipado explícito: las métricas se convierten a numérico controlando los errores, y el año "
        "queda como entero.",
        "Nulos: en las métricas críticas de negocio optamos por descartar la fila en lugar de "
        "imputar, para no inventar rendimientos que no existieron.",
        "Rangos físicos: rendimiento mayor que 0, temperatura entre -10 y 45 °C, lluvia entre 0 y "
        "12.000 mm, y pesticidas mayor o igual que 0.",
        "Consistencia referencial: una verificación automática (assert) confirma que toda clave "
        "foránea de los hechos exista en su dimensión.",
        "Conversión de unidades a la unidad de negocio: de kg/ha a ton/ha.",
    ]
)

titulo("2.3 Cumplimiento Normativo", 2)
parrafo(
    "Nos alineamos con la LOPDP (Ley Orgánica de Protección de Datos Personales del Ecuador) y con "
    "el principio de privacidad desde el diseño. En concreto: las fuentes públicas de la FAO y el "
    "Banco Mundial solo traen agregados por país, sin ningún dato personal; la fuente interna se "
    "genera anonimizada desde el inicio, con los proveedores existiendo únicamente como códigos, así "
    "que en ninguna etapa el pipeline llega a manejar PII; además, el pipeline incluye una "
    "verificación que detiene la carga si encuentra columnas de datos personales (nombre, cédula, "
    "RUC, teléfono, dirección o correo) en la fuente interna; y toda la trazabilidad de las "
    "transformaciones queda guardada en reporte_calidad.json en cada corrida."
)
doc.add_page_break()

# ============================== FASE 3 ======================================
titulo("Fase 3. Analítica Descriptiva, Diagnóstica y Visualización", 1)

titulo("3.1 Análisis Exploratorio de Datos (EDA)", 2)
est = EDA["estadisticas_descriptivas"]
parrafo(
    "Para entender las variables clave calculamos medidas de tendencia central (media y mediana) y "
    "de dispersión (desviación estándar, rango intercuartílico y coeficiente de variación), y les "
    "sumamos la asimetría y la curtosis para ver la forma de cada distribución:"
)
tabla(
    ["Variable", "Media", "Mediana", "Desv. est.", "RIQ", "CV %", "Asimetría"],
    [
        [
            v["etiqueta"],
            v["media"],
            v["mediana"],
            v["desv_estandar"],
            v["rango_intercuartilico"],
            v["coef_variacion_pct"],
            v["asimetria"],
        ]
        for v in est.values()
    ],
)
rend_global = est["rendimiento_ton_ha"]
parrafo(
    "Lo que se ve en la tabla es que en todas las variables la media queda por encima de la mediana "
    "y la asimetría es positiva, es decir, las distribuciones tienen cola hacia la derecha: hay "
    f"pocos casos con valores muy altos que estiran el promedio. El rendimiento global promedia "
    f"{rend_global['media']:.2f} ton/ha, pero la mitad de las observaciones está por debajo de "
    f"{rend_global['mediana']:.2f} ton/ha. Los pesticidas son la variable más sesgada, porque unos "
    "pocos países concentran casi todo el consumo. Como los datos no son normales, más adelante "
    "usamos también la correlación de Spearman y modelos basados en árboles, que aguantan mejor este "
    "tipo de colas largas."
)
figura("01_distribuciones.png", "Figura 1. Distribuciones de las variables clave, con la media (rojo) y la mediana (café).")
figura("02_rendimiento_ecuador.png", f"Figura 2. Rendimiento por cultivo en Ecuador ({ANIO_MIN}-{ANIO_MAX}) y evolución de los cuatro cultivos líderes.")
corr_rend = EDA["correlaciones"]["pearson"]["rendimiento_ton_ha"]
parrafo(
    "Correlación frente a causalidad. La matriz de correlaciones muestra relaciones globales débiles "
    f"entre el rendimiento y cada factor por separado (lluvia r = {corr_rend['precipitacion_mm']}, "
    f"temperatura r = {corr_rend['temperatura_media_c']}, pesticidas r = {corr_rend['pesticidas_ton']}). "
    "Esto no quiere decir que el clima no importe. Lo que ocurre es que su efecto depende del cultivo "
    "y del país, así que al mezclarlo todo la relación global se diluye. El caso de los pesticidas es "
    "un buen ejemplo de por qué correlación no es causalidad: los países que más pesticidas usan "
    "tienden a rendir más, pero el pesticida es sobre todo una señal de agricultura tecnificada "
    "(riego, semillas mejoradas, fertilización), y es esa tecnificación la que sube el rendimiento, "
    "no el pesticida por sí solo. Por eso el análisis diagnóstico se apoya además en el estudio por "
    "cultivo y en la importancia de variables del modelo (Fase 4)."
)
figura("03_correlaciones.png", "Figura 3. Correlación de Pearson y de Spearman entre el rendimiento y sus factores.")
figura("04_dispersion_factores.png", "Figura 4. Dispersión del rendimiento frente a lluvia, temperatura y pesticidas.")

titulo("3.2 Diseño de Indicadores (KPIs)", 2)
tabla(
    ["KPI", "Valor actual", "Pilar al que responde"],
    [
        [k["nombre"], k["valor"], k["objetivo_analitico"]] for k in kpis.values()
    ],
)
parrafo(
    "Los cinco KPIs cubren los cuatro pilares que definimos en la Fase 1.3: el nivel de oferta que "
    "se espera (KPI 1), el riesgo de esa oferta (KPI 2), la rentabilidad del negocio (KPI 3), las "
    "pérdidas operativas (KPI 4) y el riesgo del portafolio de compras (KPI 5). Todos salen "
    "directamente del esquema estrella y están implementados como medidas DAX en Power BI."
)
figura("05_negocio_compras.png", "Figura 5. Margen bruto anual de reventa y volumen histórico comprado por cultivo (fuente interna).")

titulo("3.3 Dashboard Estratégico y Storytelling", 2)
parrafo(
    "El dashboard lo construimos en Power BI sobre el esquema estrella (en la carpeta powerbi/ está "
    "la guía completa de armado, las medidas DAX y el tema de color). Además preparamos una versión "
    "interactiva en Streamlit para usar durante la defensa. La narrativa sigue el orden en que "
    "decidiría el comité: primero un resumen ejecutivo con los cinco KPIs, luego el diagnóstico de "
    "los factores, y al final la predicción con el plan de compras."
)
parrafo("Por qué el diseño visual es así:", negrita=True)
vinetas(
    [
        "Percepción visual. Ubicamos los KPIs en la franja superior, que es por donde entra la "
        "mirada, y ordenamos la página en una lectura tipo Z. Agrupamos por proximidad los "
        "indicadores relacionados y usamos el mismo color para el mismo cultivo entre gráficos, para "
        "que el ojo los vincule solo.",
        "Comparaciones con posición y longitud (líneas y barras), que son las que el ojo lee con más "
        "precisión. Evitamos gráficos de pastel y de área, donde los ángulos se estiman peor.",
        "Color con significado. El verde #2E7D32 es el color corporativo (agro, crecimiento, "
        "decisión positiva) y el rojo #C62828 lo reservamos solo para las alertas de reducir compra, "
        "para que con verlo ya se entienda que hay riesgo. El ámbar es para lo neutro. No pasamos de "
        "cuatro colores con significado por página.",
        "Tipografía en una sola familia sans-serif (Segoe UI), con la jerarquía dada por el tamaño y "
        "el peso, y cifras grandes en las tarjetas de KPI para que se lean de lejos en la defensa.",
    ]
)
doc.add_page_break()

# ============================== FASE 4 ======================================
titulo("Fase 4. Analítica Predictiva y Gobernanza de IA", 1)

titulo("4.1 Modelado Predictivo (Machine Learning)", 2)
parrafo(
    "Lo planteamos como un problema de regresión: estimar el rendimiento (ton/ha) a partir de seis "
    "variables predictoras, que son el año, la lluvia anual, los pesticidas, la temperatura media, "
    "el país y el tipo de cultivo. El preprocesamiento escala las variables numéricas y aplica "
    "one-hot encoding a las categóricas, todo dentro de un Pipeline de scikit-learn (así el mismo "
    "objeto sirve para entrenar y para predecir después). Partimos los datos 80/20 "
    f"({MODELO['particion']['entrenamiento']:,} filas de entrenamiento y "
    f"{MODELO['particion']['prueba']:,} de prueba) y usamos validación cruzada de 5 particiones "
    "sobre el conjunto de entrenamiento."
)
met = MODELO["metricas_por_modelo"]
tabla(
    ["Modelo", "RMSE (ton/ha)", "MAE (ton/ha)", "R² (prueba)", "R² validación cruzada (5)"],
    [
        [
            nombre,
            m["rmse_ton_ha"],
            m["mae_ton_ha"],
            m["r2"],
            f"{m['r2_validacion_cruzada_media']} ± {m['r2_validacion_cruzada_desv']}",
        ]
        for nombre, m in met.items()
    ],
)
lineal = met["Regresión Lineal"]
parrafo(
    "Para evaluar usamos las métricas que corresponden a un problema de regresión: RMSE (el error en "
    "las mismas unidades del negocio, ton/ha, que además castiga los errores grandes), MAE (el error "
    f"típico) y R² (la parte de la variabilidad que el modelo logra explicar). El ganador fue Random "
    f"Forest: explica el {rf['r2'] * 100:.1f} % de la variabilidad del rendimiento, con un RMSE de "
    f"{rf['rmse_ton_ha']} ton/ha y un MAE de {rf['mae_ton_ha']} ton/ha. La regresión lineal es más "
    f"fácil de interpretar pero se queda en un R² de {lineal['r2']:.2f}, porque la relación entre el "
    "clima y el rendimiento no es lineal y depende de interacciones (cultivo, país y clima juntos) "
    f"que los árboles capturan mejor. Que el R² de prueba y el de validación cruzada estén tan cerca "
    f"({rf['r2']:.3f} frente a {rf['r2_validacion_cruzada_media']:.3f}) nos dice que no hay "
    "sobreajuste importante."
)
figura("06_comparacion_modelos.png", "Figura 6. Comparación de los modelos: R² y RMSE en el conjunto de prueba.")
figura("07_predicho_vs_real.png", "Figura 7. Rendimiento predicho frente al real del modelo ganador (la línea marca la predicción perfecta).", 4.6)

parrafo("Pilar prescriptivo: la regla de decisión sobre la predicción.", negrita=True)
parrafo(
    "Para cada cultivo del portafolio en Ecuador predecimos el rendimiento del siguiente ciclo bajo "
    "un escenario climático base (el promedio de los últimos tres años) y calculamos un índice de "
    "oferta, que es el rendimiento predicho dividido entre el promedio histórico. La regla es "
    "sencilla: si el índice es mayor o igual a 1,05 conviene aumentar la compra un 15 %, porque va a "
    "haber oferta abundante y los precios tienden a bajar; si queda entre 0,95 y 1,05 se mantiene; y "
    "si es menor que 0,95 se reduce un 15 % y se aseguran contratos por adelantado, porque el riesgo "
    "es de escasez y sobreprecio."
)
tabla(
    ["Cultivo", "Predicho (ton/ha)", "Histórico (ton/ha)", "Índice", "Decisión"],
    [
        [
            f["cultivo"],
            f["rendimiento_predicho_ton_ha"],
            f["rendimiento_promedio_historico_ton_ha"],
            f["indice_oferta"],
            f["decision"],
        ]
        for _, f in RECOMENDACIONES.iterrows()
    ],
)

titulo("4.2 Ética de la IA y Gobernanza de Datos", 2)
parrafo("Sesgos algorítmicos: nuestro análisis crítico.", negrita=True)
sesgo = MODELO["sesgo_representacion"]
pais_top, filas_top = next(iter(sesgo["paises_mas_representados"].items()))
pais_min, filas_min = list(sesgo["paises_menos_representados"].items())[-1]
errores = MODELO["error_por_grupo"]["rmse_por_cultivo"]
cultivo_peor = max(errores, key=errores.get)
cultivo_mejor = min(errores, key=errores.get)
filas_cultivo = sesgo["cultivos"]
cultivo_menos_repr = min(filas_cultivo, key=filas_cultivo.get)
cultivo_mas_repr = max(filas_cultivo, key=filas_cultivo.get)
vinetas(
    [
        "Sesgo de representación. El dataset tiene muchas más filas de países con estadísticas "
        f"completas ({pais_top}: {filas_top:,} filas) que de otros con series cortas ({pais_min}: "
        f"{filas_min}). Con los cultivos pasa parecido: los de alcance global ({cultivo_mas_repr}: "
        f"{filas_cultivo[cultivo_mas_repr]:,} filas) pesan más que los más regionales "
        f"({cultivo_menos_repr}: {filas_cultivo[cultivo_menos_repr]:,}). Ecuador tiene "
        f"{sesgo['filas_ecuador']} filas, una representación intermedia que nos parece razonable.",
        "El riesgo de esto es que el modelo aprenda mejor a los grupos dominantes y prediga peor, es "
        "decir, 'discrimine' en términos estadísticos, a los países o cultivos con pocos datos.",
        f"Para no quedarnos en la sospecha lo medimos. El RMSE para Ecuador es de "
        f"{MODELO['error_por_grupo']['rmse_ecuador']} ton/ha, mejor que el promedio global de "
        f"{rf['rmse_ton_ha']}, así que el sesgo no perjudica a nuestro mercado objetivo. Por cultivo "
        f"el error va de {errores[cultivo_mejor]} ({cultivo_mejor}) a {errores[cultivo_peor]} "
        f"({cultivo_peor}); revisándolo, ese error se explica más por la escala y la varianza del "
        "rendimiento de cada cultivo (la papa va de 5 a 50 ton/ha según el país, mientras que la "
        "soya se mueve en un rango estrecho) que por la representación. Aun así, las recomendaciones "
        "de los cultivos con más error se entregan al negocio con una advertencia de mayor "
        "incertidumbre.",
        f"Sesgo temporal. El histórico llega hasta {ANIO_MAX}, porque FAOSTAT publica con uno o dos "
        "años de rezago. El pipeline automatizado permite reentrenar en cuanto salga la "
        "actualización oficial, antes de cada campaña.",
    ]
)
parrafo("IA explicable (XAI):", negrita=True)
xai = MODELO["xai_importancia_permutacion"]
parrafo(
    "Para que el comité pueda confiar en el modelo calculamos la importancia de cada variable por "
    "permutación, que mide cuánto baja el R² cuando se desordena esa variable. El resultado tiene "
    "sentido para el negocio: el tipo de cultivo es de lejos lo que más pesa (una papa rinde en "
    "toneladas muchísimo más que un trigo), seguido del país, que funciona como resumen del suelo, "
    "la tecnología y las variedades de cada lugar. Entre los factores que sí se pueden monitorear o "
    f"ajustar, los pesticidas ({xai['Pesticidas (ton)']:.2f}), la temperatura "
    f"({xai['Temperatura media (°C)']:.2f}) y la lluvia ({xai['Lluvia anual (mm)']:.2f}) tienen un "
    "peso parecido. Dicho de forma simple: el modelo primero identifica el cultivo y el país, y "
    "luego ajusta la estimación con el clima y el manejo del año. Es parecido a cómo razonaría un "
    f"agrónomo con experiencia, solo que con la memoria de {N_FILAS:,} cosechas encima."
)
figura("08_importancia_variables.png", "Figura 8. Importancia de variables por permutación (XAI) del modelo ganador.")
parrafo(
    "Gobernanza del modelo. El pipeline completo (preprocesamiento y modelo) se guarda como un solo "
    "archivo versionado (modelo_rendimiento.joblib); las predicciones siempre se comunican con su "
    f"margen de error (±{rf['rmse_ton_ha']} ton/ha); y la decisión final de compra la sigue tomando "
    "el comité, con el modelo como apoyo. Esta supervisión humana la dejamos definida desde la Fase 1."
)
doc.add_page_break()

# ========================== CONCLUSIONES ====================================
titulo("Conclusiones", 1)
vinetas(
    [
        "Cerramos el ciclo analítico completo: integramos tres fuentes (una de ellas una API "
        "pública), documentamos un esquema estrella, hicimos un EDA con base estadística, definimos "
        f"cinco KPIs, entrenamos un modelo predictivo con R² de {rf['r2']:.3f} y armamos una regla "
        "prescriptiva que convierte esa predicción en decisiones de compra concretas.",
        "El valor para el negocio es directo: anticipar la oferta permite comprar más cuando se "
        "espera abundancia (y mejores precios) y asegurar contratos cuando se espera escasez, lo que "
        "reduce a la vez el sobrestock y el quiebre de stock.",
        "También somos honestos con las limitaciones: el error es mayor en los cultivos con pocos "
        f"datos, FAOSTAT publica con rezago (el histórico llega hasta {ANIO_MAX}) y las compras "
        "internas son simuladas. Cada una de esas limitaciones tiene su mitigación documentada.",
        "Como trabajo futuro queda reentrenar el modelo con cada actualización de FAOSTAT, sumar "
        "precios de mercado reales para pasar de una regla de decisión a una optimización de margen, "
        "y bajar el análisis a nivel de provincia cuando existan datos subnacionales.",
    ]
)

RUTA_SALIDA.parent.mkdir(exist_ok=True)
doc.save(RUTA_SALIDA)
print(f"OK — informe generado en {RUTA_SALIDA}")
